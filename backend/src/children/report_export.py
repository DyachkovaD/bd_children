"""
Генерация отчётов по учащимся в форматах xlsx и docx.
"""
from io import BytesIO
from datetime import datetime

# Поля отчёта (ключ в данных API -> заголовок в файле)
REPORT_FIELDS = [
    ('last_name', 'Фамилия'),
    ('first_name', 'Имя'),
    ('patronymic', 'Отчество'),
    ('birthday', 'Дата рождения'),
    ('age', 'Возраст'),
    ('education_class', 'Класс'),
    ('school_name', 'Школа'),
    ('address', 'Адрес'),
    ('health_status', 'Состояние здоровья'),
    ('family_status', 'Статус семьи'),
    ('note', 'Примечание'),
]

REPORT_FIELD_KEYS = [k for k, _ in REPORT_FIELDS]


def _format_cell_value(value, key):
    if value is None or value == '':
        return ''
    if key == 'birthday' and value:
        try:
            if isinstance(value, str):
                dt = datetime.fromisoformat(value.replace('Z', '+00:00'))
            else:
                dt = value
            return dt.strftime('%d.%m.%Y')
        except (ValueError, TypeError):
            return str(value)
    if key == 'age' and value is not None and value != '':
        return str(value)
    return '' if value is None or value == '' else str(value)


def build_xlsx(rows_data, selected_headers):
    """Строит xlsx-файл. selected_headers — список кортежей (key, label)."""
    import openpyxl
    from openpyxl.styles import Font, Alignment
    from openpyxl.utils import get_column_letter

    WIDTH_MULTIPLIER = 1.5
    HEIGHT_MULTIPLIER = 1.0  # уменьшенный межстрочный интервал
    DEFAULT_COLUMN_WIDTH = 8.43
    DEFAULT_ROW_HEIGHT = 15
    CENTER_HORIZONTAL_KEYS = ('education_class', 'age', 'birthday')

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Учащиеся'
    headers = [label for _, label in selected_headers]
    ws.append(headers)
    for col in range(1, len(headers) + 1):
        ws.cell(row=1, column=col).font = Font(bold=True)
    for row_data in rows_data:
        row = []
        for key, _ in selected_headers:
            row.append(_format_cell_value(row_data.get(key), key))
        ws.append(row)

    # Ширина колонок в 1.5 раза больше; Школа, Адрес, Состояние здоровья, Статус семьи, Примечание — в 2 раза шире
    col_width = DEFAULT_COLUMN_WIDTH * WIDTH_MULTIPLIER
    WIDE_COLUMN_KEYS = ('school_name', 'address', 'health_status', 'family_status', 'note')
    col_widths = []
    for col_idx, (key, _) in enumerate(selected_headers):
        col_num = col_idx + 1
        width = col_width * 2 if key in WIDE_COLUMN_KEYS else col_width
        ws.column_dimensions[get_column_letter(col_num)].width = width
        col_widths.append(width)

    # Автоподбор высоты строки по содержимому для каждой строки
    import math
    chars_per_unit = 0.9
    points_per_line = 15
    for row in range(1, ws.max_row + 1):
        max_lines = 1
        for col_idx in range(len(selected_headers)):
            cell = ws.cell(row=row, column=col_idx + 1)
            text = (cell.value or '')
            if isinstance(text, (int, float)):
                text = str(text)
            w = col_widths[col_idx]
            chars_fit = max(1, w * chars_per_unit)
            max_lines = max(max_lines, math.ceil(len(text) / chars_fit))
        ws.row_dimensions[row].height = max(DEFAULT_ROW_HEIGHT, points_per_line * max_lines)

    # Вертикальное выравнивание по центру для всех ячеек; горизонтальное по центру для Класс, Возраст, Дата рождения
    for col_idx, (key, _) in enumerate(selected_headers):
        col_num = col_idx + 1
        h_center = key in CENTER_HORIZONTAL_KEYS
        align = Alignment(
            horizontal='center' if h_center else 'general',
            vertical='center',
            wrap_text=True,
        )
        for row in range(1, ws.max_row + 1):
            ws.cell(row=row, column=col_num).alignment = align

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


def build_docx(rows_data, selected_headers):
    """Строит docx-файл. selected_headers — список кортежей (key, label)."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    doc = Document()
    # Узкие поля страницы
    section = doc.sections[0]
    narrow = Inches(0.5)
    section.top_margin = narrow
    section.bottom_margin = narrow
    section.left_margin = narrow
    section.right_margin = narrow
    doc.add_heading('Отчёт по учащимся', level=0)
    headers = [label for _, label in selected_headers]
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, label in enumerate(headers):
        hdr_cells[i].text = label
        for p in hdr_cells[i].paragraphs:
            if p.runs:
                p.runs[0].bold = True
    for row_idx, row_data in enumerate(rows_data):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, (key, _) in enumerate(selected_headers):
            row_cells[col_idx].text = _format_cell_value(row_data.get(key), key)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def get_selected_headers(fields_param):
    """
    Возвращает список кортежей (key, label) для выбранных полей.
    fields_param — строка через запятую (ключи) или None/пустая — все поля.
    """
    if not fields_param or not fields_param.strip():
        return list(REPORT_FIELDS)
    keys = [k.strip() for k in fields_param.split(',') if k.strip()]
    key_to_label = dict(REPORT_FIELDS)
    return [(k, key_to_label[k]) for k in keys if k in key_to_label]
